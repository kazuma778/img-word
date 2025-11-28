from flask import Flask, render_template, request, send_file, redirect, url_for, send_from_directory, after_this_request, jsonify, flash, abort, render_template_string
import os
import zipfile
from docx import Document
from PIL import Image
import tempfile
import shutil
try:
    import pythoncom
    import win32com.client as win32
except ImportError:
    pythoncom = None
    win32 = None
import fitz  # PyMuPDF
import uuid
from werkzeug.utils import secure_filename
import time
import threading
import io
import docx
from datetime import datetime
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer
from docx import Document as DocxDocument
from markupsafe import Markup
import re, os
import json
import requests # Added for image upscaling
import logging

# Setup file logging (logs to log.txt, cleared on restart)
from setup_logging import setup_file_logging
logger = setup_file_logging()

# --- Konfigurasi Upscaling ---
# Cookie otentikasi Anda.
COOKIE_VALUE = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpZCI6MzIwOTY1LCJ1c2VybmFtZSI6ImNvZmZlZXNvY2lhbEAyMjAwZnJlZWZvbnRzLmNvbSIsInBhc3N3b3JkIjoiMzQ0MTIzRkRBOEJENDM4NkVFODg5MzYzMzE2QTA0Q0EiLCJpYXQiOjE3NjIzMzkwNDF9.05Nar5hAACa6qGpcmngmQBucS49765nLDc8Uo5mcqkw"

# Skala pembesaran: "200" untuk 2x, "400" untuk 4x.
# Untuk saat ini, kita akan menggunakan 4x secara default.
UPSCALE_FACTOR = "400"

# Header yang akan digunakan untuk semua permintaan.
HEADERS = {
    "Cookie": f"jwt={COOKIE_VALUE}",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
    "Accept": "application/json",
    "Origin": "https://imgupscaler.com",
    "Referer": "https://imgupscaler.com/"
}

# URL API
UPLOAD_URL = "https://get1.imglarger.com/api/UpscalerNew/UploadNew"
STATUS_URL = "https://get1.imglarger.com/api/UpscalerNew/CheckStatusNew"

def upload_image(file_path):
    """
    Mengunggah gambar dan mengembalikan ID tugas (code) dan nilai skala.
    """
    if not os.path.exists(file_path):
        print(f"Error: File tidak ditemukan di {file_path}")
        return None, None

    print(f"Mengunggah {os.path.basename(file_path)}...")
    
    scale_map = {"200": "1", "400": "4"}
    scale_value = scale_map.get(UPSCALE_FACTOR, "4")

    with open(file_path, 'rb') as f:
        files = {'myfile': (os.path.basename(file_path), f, 'image/jpeg')}
        data = {'scaleRadio': scale_value}
        
        try:
            response = requests.post(UPLOAD_URL, headers=HEADERS, files=files, data=data)
            response.raise_for_status()
            upload_data = response.json()
            
            if upload_data.get("msg") == "Success":
                task_id = upload_data.get("data", {}).get("code")
                if task_id:
                    print(f"Upload berhasil. Task ID (code): {task_id}")
                    return task_id, scale_value
                else:
                    print("Upload berhasil tetapi tidak menemukan Task ID (code) dalam respons.")
                    print(upload_data)
                    return None, None
            else:
                print("Upload gagal. Respons dari server:")
                print(upload_data)
                return None, None

        except requests.exceptions.RequestException as e:
            print(f"Terjadi error saat mengunggah: {e}")
            return None, None
        except ValueError:
            print("Gagal mem-parse respons dari server. Respons mentah:")
            print(response.text)
            return None, None

def check_status(task_id, scale_value):
    """
    Memeriksa status proses upscaling secara berkala hingga selesai.
    """
    print(f"Memeriksa status untuk Task ID (code): {task_id}")
    
    POLLING_INTERVAL = 10
    MAX_WAIT_TIME = 300
    start_time = time.time()

    while True:
        elapsed_time = time.time() - start_time
        if elapsed_time > MAX_WAIT_TIME:
            print(f"\nBatas waktu {MAX_WAIT_TIME} detik terlampaui. Proses dibatalkan.")
            return None

        try:
            data = {"code": task_id, "scaleRadio": scale_value}
            response = requests.post(STATUS_URL, headers=HEADERS, json=data)
            response.raise_for_status()
            status_data = response.json()
            
            if status_data.get("msg") == "Success":
                status_info = status_data.get("data", {})
                status = status_info.get("status")
                print(f"\rStatus saat ini: {status} (menunggu {int(elapsed_time)} detik)...")

                if status == "success":
                    download_urls = status_info.get("downloadUrls", [])
                    if download_urls:
                        download_url = download_urls[0]
                        print(f"\nProses selesai! URL download: {download_url}")
                        return download_url
                    else:
                        print("\nProses berhasil tetapi tidak menemukan URL download.")
                        return None
                elif status == "fail":
                    print("\nProses upscaling gagal.")
                    return None
                elif status == "waiting":
                    pass
                else:
                    print(f"\nStatus tidak dikenali: {status}")

            else:
                print(f"\nGagal memeriksa status. Respons: {status_data}")
                return None

            time.sleep(POLLING_INTERVAL)

        except requests.exceptions.RequestException as e:
            print(f"\nTerjadi error saat memeriksa status: {e}")
            return None
        except ValueError:
            print(f"\nGagal mem-parse respons status. Respons mentah: {response.text}")
            return None

def download_result(url, output_path):
    """
    Mengunduh gambar hasil upscale dari URL yang diberikan.
    """
    print(f"Mengunduh gambar dari {url}...")
    try:
        download_headers = {"User-Agent": HEADERS["User-Agent"]}
        response = requests.get(url, headers=download_headers, stream=True)
        response.raise_for_status()
        
        with open(output_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        print(f"Gambar berhasil disimpan di {output_path}")
        return True

    except requests.exceptions.RequestException as e:
        print(f"Gagal mengunduh gambar: {e}")
        return False

app = Flask(__name__)
app.secret_key = str(uuid.uuid4())

# Configuration
if os.name == 'nt':
    # Windows (Local Development)
    app.config['UPLOAD_FOLDER'] = 'uploads'
    app.config['PROCESSED_FOLDER'] = 'processed'
else:
    # Linux (Vercel/Production) - Use /tmp for writable storage
    app.config['UPLOAD_FOLDER'] = '/tmp/uploads'
    app.config['PROCESSED_FOLDER'] = '/tmp/processed'
app.config['MAX_FILES'] = 5  # Maksimal 5 file di processed
app.config['MAX_UPLOAD_FILES'] = 500  # Maksimal 5 file di uploads
app.config['ALLOWED_EXTENSIONS'] = {'docx', 'pdf', 'doc', 'jpg', 'jpeg', 'png', 'rtf'}  # Tambah RTF

# Dictionary to store task progress (REMOVED - Replaced by File-Based Storage)
# tasks = {}

# Task Storage Configuration
if os.name == 'nt':
    TASKS_DIR = os.path.join(os.getcwd(), 'processed', 'tasks')
else:
    TASKS_DIR = os.path.join('/tmp', 'processed', 'tasks')

os.makedirs(TASKS_DIR, exist_ok=True)

def save_task(task_id, data):
    """Save task status to JSON file"""
    try:
        file_path = os.path.join(TASKS_DIR, f"{task_id}.json")
        with open(file_path, 'w') as f:
            json.dump(data, f)
    except Exception as e:
        print(f"Error saving task {task_id}: {e}")

def get_task(task_id):
    """Get task status from JSON file"""
    try:
        file_path = os.path.join(TASKS_DIR, f"{task_id}.json")
        if os.path.exists(file_path):
            with open(file_path, 'r') as f:
                return json.load(f)
        return None
    except Exception as e:
        print(f"Error reading task {task_id}: {e}")
        return None

def update_task(task_id, **kwargs):
    """Update specific fields in task"""
    data = get_task(task_id)
    if data:
        data.update(kwargs)
        save_task(task_id, data)

# Create necessary directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['PROCESSED_FOLDER'], exist_ok=True)

def allowed_file(filename, allowed_extensions=None):
    if allowed_extensions is None:
        allowed_extensions = app.config['ALLOWED_EXTENSIONS']
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def cleanup_processed_folder():
    """Hapus item lama di folder processed jika melebihi batas 5 (termasuk file & folder)."""
    processed_dir = app.config['PROCESSED_FOLDER']
    max_items = 5  # Keep the 5 most recent items

    try:
        # Get all items (files and folders) in the directory
        all_item_names = os.listdir(processed_dir)
        
        # Create a list of (path, mod_time) tuples
        items_with_time = []
        for item_name in all_item_names:
            item_path = os.path.join(processed_dir, item_name)
            try:
                items_with_time.append((item_path, os.path.getmtime(item_path)))
            except FileNotFoundError:
                # Item might be deleted by another process between listdir and getmtime
                continue

        # Sort items by modification time, newest first
        items_with_time.sort(key=lambda x: x[1], reverse=True)
        
        # Identify items to delete (all items after the 5th one)
        items_to_delete = items_with_time[max_items:]
        
        # Delete the old items
        if items_to_delete:
            print(f"Cleanup: Ditemukan {len(items_with_time)} item, akan dihapus {len(items_to_delete)} item terlama.")
        
        for path_to_delete, mod_time in items_to_delete:
            try:
                if os.path.isfile(path_to_delete):
                    os.remove(path_to_delete)
                    print(f"Menghapus file processed lama: {os.path.basename(path_to_delete)}")
                elif os.path.isdir(path_to_delete):
                    shutil.rmtree(path_to_delete)
                    print(f"Menghapus folder processed lama: {os.path.basename(path_to_delete)}")
            except Exception as e:
                print(f"Gagal menghapus item lama '{os.path.basename(path_to_delete)}': {e}")
                
    except Exception as e:
        print(f"Error pada saat pembersihan folder processed: {e}")

def cleanup_uploads_folder():
    """Hapus item lama di folder uploads jika melebihi batas maksimal (termasuk file & folder)."""
    uploads_dir = app.config['UPLOAD_FOLDER']
    max_items = app.config['MAX_UPLOAD_FILES']

    try:
        # Get all items (files and folders) in the directory
        all_item_names = os.listdir(uploads_dir)
        
        # Create a list of (path, mod_time) tuples
        items_with_time = []
        for item_name in all_item_names:
            item_path = os.path.join(uploads_dir, item_name)
            try:
                items_with_time.append((item_path, os.path.getmtime(item_path)))
            except FileNotFoundError:
                continue

        # Sort items by modification time, newest first
        items_with_time.sort(key=lambda x: x[1], reverse=True)
        
        # Identify items to delete
        items_to_delete = items_with_time[max_items:]
        
        if items_to_delete:
            print(f"Cleanup Uploads: Ditemukan {len(items_with_time)} item, akan dihapus {len(items_to_delete)} item terlama.")

        for path_to_delete, mod_time in items_to_delete:
            try:
                if os.path.isfile(path_to_delete):
                    os.remove(path_to_delete)
                    # print(f"Menghapus file upload lama: {os.path.basename(path_to_delete)}")
                elif os.path.isdir(path_to_delete):
                    shutil.rmtree(path_to_delete)
                    # print(f"Menghapus folder upload lama: {os.path.basename(path_to_delete)}")
            except Exception as e:
                print(f"Gagal menghapus item upload lama '{os.path.basename(path_to_delete)}': {e}")
                
    except Exception as e:
        print(f"Error pada saat pembersihan folder uploads: {e}")

def cleanup_all_folders():
    """Bersihkan semua folder sementara"""
    cleanup_uploads_folder()
    cleanup_processed_folder()
    cleanup_tasks_folder()

def cleanup_tasks_folder():
    """Hapus file task lama (> 1 jam)"""
    try:
        current_time = time.time()
        for filename in os.listdir(TASKS_DIR):
            file_path = os.path.join(TASKS_DIR, filename)
            # Hapus jika lebih dari 1 jam
            if os.path.getmtime(file_path) < current_time - 3600:
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Gagal menghapus task lama {filename}: {e}")
    except Exception as e:
        print(f"Error cleanup tasks: {e}")

def periodic_cleanup():
    """Cleanup berkala setiap 30 menit"""
    while True:
        time.sleep(1800)  # 30 menit = 1800 detik
        try:
            print("Melakukan cleanup berkala...")
            cleanup_all_folders()
            print("Cleanup berkala selesai")
        except Exception as e:
            print(f"Error pada cleanup berkala: {e}")

def standardize_image_format(image_path, output_dir, image_counter, target_format='JPEG'):
    """Standardisasi format gambar dan beri nama berurutan"""
    try:
        with Image.open(image_path) as img:
            # Konversi ke RGB jika diperlukan untuk JPEG
            if target_format.upper() == 'JPEG' and img.mode in ('RGBA', 'LA', 'P'):
                # Buat background putih untuk gambar transparan
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            elif target_format.upper() == 'PNG' and img.mode not in ('RGBA', 'RGB', 'L'):
                img = img.convert('RGBA')

            # Tentukan extension berdasarkan format
            ext = '.jpg' if target_format.upper() == 'JPEG' else '.png'
            
            # Buat nama file berurutan
            standardized_name = f"image_{image_counter:03d}{ext}"
            standardized_path = os.path.join(output_dir, standardized_name)
            
            # Simpan dengan format yang distandarisasi
            img.save(standardized_path, format=target_format.upper(), quality=100 if target_format.upper() == 'JPEG' else None)
            
            return standardized_path, standardized_name
    except Exception as e:
        print(f"Error standardisasi format gambar {image_path}: {e}")
        return None, None

def extract_images_from_pdf_sequential(pdf_path, temp_dir, task_id=None, file_index=0, total_files=1):
    """Ekstrak gambar dari PDF dengan urutan yang benar berdasarkan halaman dan posisi"""
    try:
        doc = fitz.open(pdf_path)
        all_images_data = []  # List untuk menyimpan semua data gambar dengan urutan
        total_pages = len(doc)

        # Kumpulkan semua gambar dari semua halaman dengan informasi urutan
        for page_num in range(total_pages):
            if task_id and (page_num % 5 == 0 or page_num == total_pages - 1):
                progress = 20 + (file_index * 30 / total_files) + ((page_num / total_pages) * 15 / total_files)
                update_task_progress(task_id, progress, f"Scanning PDF {file_index+1}/{total_files}: page {page_num+1}/{total_pages}")

            page = doc.load_page(page_num)
            page_images = page.get_images(full=True)

            # Untuk setiap gambar di halaman ini
            for img_index, img in enumerate(page_images):
                try:
                    # Dapatkan posisi gambar untuk pengurutan
                    image_instances = page.get_image_rects(img[0])
                    y_position = 0
                    if image_instances:
                        y_position = image_instances.y0

                    # Simpan data gambar dengan informasi urutan
                    all_images_data.append({
                        'page_num': page_num,
                        'img_index': img_index,
                        'y_position': y_position,
                        'img_data': img,
                        'sort_key': (page_num, y_position, img_index)  # Key untuk sorting
                    })
                except Exception as e:
                    print(f"Error mengambil posisi gambar {img_index} dari halaman {page_num}: {e}")
                    # Tetap simpan gambar meskipun tidak bisa mendapat posisi
                    all_images_data.append({
                        'page_num': page_num,
                        'img_index': img_index,
                        'y_position': 0,
                        'img_data': img,
                        'sort_key': (page_num, 0, img_index)
                    })

        # Urutkan semua gambar berdasarkan halaman, kemudian posisi Y
        all_images_data.sort(key=lambda x: x['sort_key'])

        # Ekstrak gambar sesuai urutan yang benar
        extracted_images = []
        total_images = len(all_images_data)
        
        for sequential_index, img_data in enumerate(all_images_data):
            if task_id and (sequential_index % 10 == 0 or sequential_index == total_images - 1):
                progress = 35 + (file_index * 30 / total_files) + ((sequential_index / total_images) * 15 / total_files)
                update_task_progress(task_id, progress, f"Extracting image {sequential_index+1}/{total_images} from PDF {file_index+1}/{total_files}")

            try:
                img = img_data['img_data']
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_ext = base_image["ext"].lower()

                # Buat nama file sementara dengan urutan yang benar
                temp_image_name = f"temp_seq_{sequential_index+1:03d}.{image_ext}"
                temp_image_path = os.path.join(temp_dir, temp_image_name)

                with open(temp_image_path, "wb") as img_file:
                    img_file.write(base_image["image"])

                # Standardisasi format dengan nomor urut yang benar
                final_counter = sequential_index + 1
                standardized_path, standardized_name = standardize_image_format(
                    temp_image_path, temp_dir, final_counter, 'JPEG'
                )

                if standardized_path:
                    extracted_images.append((standardized_path, standardized_name, final_counter))

                # Hapus file temporary
                safe_remove_file(temp_image_path)

            except Exception as e:
                print(f"Gagal ekstrak gambar {sequential_index}: {e}")

        doc.close()
        return extracted_images

    except Exception as e:
        print(f"Error membuka PDF {pdf_path}: {e}")
        return []

def extract_images_from_docx_sequential(file_path, output_dir, task_id=None, file_index=0, total_files=1):
    """Ekstrak gambar dari DOCX dengan urutan yang benar berdasarkan urutan kemunculan dalam dokumen secara linear"""
    try:
        doc = Document(file_path)
        
        # Kumpulkan gambar dengan urutan berdasarkan posisi dalam struktur dokumen
        ordered_images = []
        
        # Buka dokumen sebagai ZIP untuk mengakses struktur XML
        import zipfile
        import re
        
        with zipfile.ZipFile(file_path, 'r') as docx_zip:
            # Baca document.xml untuk mendapatkan urutan linear gambar
            document_xml = docx_zip.read('word/document.xml').decode('utf-8')
            
            # Pattern untuk mencari referensi gambar dalam urutan kemunculan
            # Cari semua r:embed dalam urutan kemunculan di document.xml
            embed_pattern = r'r:embed="([^"]+)"'
            embed_matches = re.findall(embed_pattern, document_xml)
            
            # Buat mapping urutan berdasarkan kemunculan linear dalam document.xml
            embed_order = {}
            for idx, embed_id in enumerate(embed_matches):
                if embed_id not in embed_order:  # Hanya ambil kemunculan pertama untuk urutan
                    embed_order[embed_id] = idx

        # Kumpulkan semua relasi gambar dengan urutan linear
        image_relations_with_order = []
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                # Tentukan urutan berdasarkan kemunculan linear dalam document.xml
                linear_order = embed_order.get(rel_id, 999999)  # Default tinggi untuk yang tidak ditemukan
                image_relations_with_order.append({
                    'order': linear_order,
                    'rel': rel,
                    'rel_id': rel_id
                })

        # Urutkan berdasarkan urutan kemunculan linear dalam dokumen (meniru urutan halaman)
        image_relations_with_order.sort(key=lambda x: x['order'])

        extracted_images = []
        total_images = len(image_relations_with_order)
        
        for sequential_index, img_data in enumerate(image_relations_with_order):
            if task_id and (sequential_index % 10 == 0 or sequential_index == total_images - 1):
                progress = 20 + (file_index * 30 / total_files) + ((sequential_index / total_images) * 30 / total_files)
                update_task_progress(task_id, progress, f"Extracting image {sequential_index+1}/{total_images} from document {file_index+1}/{total_files}")

            try:
                rel = img_data['rel']
                
                # Ekstrak data gambar
                img_blob = rel.target_part.blob
                
                # Tentukan format dari nama file asli
                original_name = os.path.basename(rel.target_ref)
                original_ext = os.path.splitext(original_name)[1].lower()
                
                # Buat file temporary dengan urutan yang benar
                temp_name = f"temp_seq_{sequential_index+1:03d}{original_ext}"
                temp_path = os.path.join(output_dir, temp_name)
                
                with open(temp_path, 'wb') as f:
                    f.write(img_blob)

                # Standardisasi format dengan nomor urut yang benar
                final_counter = sequential_index + 1
                standardized_path, standardized_name = standardize_image_format(
                    temp_path, output_dir, final_counter, 'JPEG'
                )

                if standardized_path:
                    extracted_images.append((standardized_path, standardized_name, final_counter))

                # Hapus file temporary
                safe_remove_file(temp_path)

            except Exception as e:
                print(f"Gagal ekstrak gambar {sequential_index}: {e}")

        return extracted_images

    except zipfile.BadZipFile as e:
        print(f"File DOCX rusak {file_path}: {e}")
        return []
    except Exception as e:
        print(f"Error ekstrak gambar dari DOCX {file_path}: {e}")
        return []

def convert_images_to_grayscale_ordered(images_data, task_id=None, progress_start=60, progress_range=30):
    """Konversi gambar ke grayscale dengan mempertahankan urutan"""
    grayscale_images = []
    total_images = len(images_data)
    
    # Urutkan berdasarkan counter untuk memastikan urutan yang benar
    images_data.sort(key=lambda x: x[2])  # Sort by counter
    
    for i, (image_path, original_name, counter) in enumerate(images_data):
        try:
            if task_id and (i % 10 == 0 or i == total_images - 1):
                progress = progress_start + (i / total_images) * progress_range
                update_task_progress(task_id, progress, f"Converting image {i+1}/{total_images} to grayscale")

            if not os.path.exists(image_path):
                continue

            with Image.open(image_path) as img:
                # Konversi ke grayscale
                grayscale_img = img.convert('L')
                
                # Buat nama file grayscale dengan urutan yang sama
                base_name = os.path.splitext(original_name)[0]
                grayscale_name = f"{base_name}_grayscale.jpg"
                grayscale_path = os.path.join(os.path.dirname(image_path), grayscale_name)
                
                # Simpan gambar grayscale
                grayscale_img.save(grayscale_path, 'JPEG', quality=100)
                grayscale_images.append((grayscale_path, grayscale_name, counter))

                # Hapus file asli setelah konversi
                safe_remove_file(image_path)

        except Exception as e:
            print(f"Error konversi grayscale {image_path}: {e}")

    return grayscale_images

def doc_to_docx(doc_path):
    """Konversi DOC ke DOCX dengan penanganan error yang lebih baik - FIXED"""
    if pythoncom is None or win32 is None:
        print("Windows COM libraries not available. Skipping DOC conversion.")
        return None

    pythoncom.CoInitialize()
    word = None
    try:
        # Pastikan menggunakan absolute path
        doc_abs_path = os.path.abspath(doc_path)
        print(f"DOC absolute path: {doc_abs_path}")
        
        # Pastikan file DOC ada
        if not os.path.exists(doc_abs_path):
            print(f"File DOC tidak ditemukan: {doc_abs_path}")
            return None
            
        # Buat absolute path untuk output DOCX
        docx_abs_path = os.path.abspath(os.path.splitext(doc_path)[0] + '.docx')
        print(f"DOCX absolute path: {docx_abs_path}")
        
        # Pastikan direktori output exists
        output_dir = os.path.dirname(docx_abs_path)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False  # Pastikan Word tidak muncul
        word.DisplayAlerts = False  # Nonaktifkan alert dialog
        
        time.sleep(1)
        
        # Buka file DOC dengan absolute path
        print(f"Membuka file DOC: {doc_abs_path}")
        doc = word.Documents.Open(doc_abs_path)
        
        # Simpan sebagai DOCX dengan absolute path
        print(f"Menyimpan sebagai DOCX: {docx_abs_path}")
        doc.SaveAs(docx_abs_path, FileFormat=16)  # FileFormat 16 = DOCX
        doc.Close()
        
        print(f"Konversi berhasil: {docx_abs_path}")
        return docx_abs_path
        
    except Exception as e:
        print(f"Error konversi DOC ke DOCX: {e}")
        print(f"DOC path: {doc_path}")
        return None
    finally:
        if word:
            try:
                word.Quit()
                time.sleep(1)
            except Exception as e:
                print(f"Error closing Word: {e}")

def rtf_to_docx(rtf_path):
    """Konversi RTF ke DOCX dengan penanganan error yang lebih baik - FIXED"""
    if pythoncom is None or win32 is None:
        print("Windows COM libraries not available. Skipping RTF conversion.")
        return None

    pythoncom.CoInitialize()
    word = None
    try:
        # Pastikan menggunakan absolute path
        rtf_abs_path = os.path.abspath(rtf_path)
        print(f"RTF absolute path: {rtf_abs_path}")
        
        # Pastikan file RTF ada
        if not os.path.exists(rtf_abs_path):
            print(f"File RTF tidak ditemukan: {rtf_abs_path}")
            return None
            
        # Buat absolute path untuk output DOCX
        docx_abs_path = os.path.abspath(os.path.splitext(rtf_path)[0] + '.docx')
        print(f"DOCX absolute path: {docx_abs_path}")
        
        # Pastikan direktori output exists
        output_dir = os.path.dirname(docx_abs_path)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False  # Pastikan Word tidak muncul
        word.DisplayAlerts = False  # Nonaktifkan alert dialog
        
        time.sleep(1)
        
        # Buka file RTF dengan absolute path
        print(f"Membuka file RTF: {rtf_abs_path}")
        doc = word.Documents.Open(rtf_abs_path)
        
        # Simpan sebagai DOCX dengan absolute path
        print(f"Menyimpan sebagai DOCX: {docx_abs_path}")
        doc.SaveAs(docx_abs_path, FileFormat=16)  # FileFormat 16 = DOCX
        doc.Close()
        
        print(f"Konversi berhasil: {docx_abs_path}")
        return docx_abs_path
        
    except Exception as e:
        print(f"Error konversi RTF ke DOCX: {e}")
        print(f"RTF path: {rtf_path}")
        return None
    finally:
        if word:
            try:
                word.Quit()
                time.sleep(1)
            except Exception as e:
                print(f"Error closing Word: {e}")

def safe_remove_file(file_path, retry=3, delay=0.5):
    """Hapus file dengan aman dengan beberapa percobaan"""
    for i in range(retry):
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
            return True
        except Exception as e:
            print(f"Gagal menghapus file (percobaan {i+1}): {e}")
            time.sleep(delay)
    return False

def update_task_progress(task_id, progress, message=None):
    """Update progress for a specific task"""
    update_task(task_id, progress=min(progress, 99), message=message)
    print(f"Task {task_id}: {progress:.1f}% - {message}")

def complete_task(task_id, filename):
    """Mark a task as completed"""
    update_task(task_id, status='completed', progress=100, download_filename=filename)
    print(f"Task {task_id} completed: {filename}")

def fail_task(task_id, error_message):
    """Mark a task as failed"""
    update_task(task_id, status='failed', message=error_message)
    print(f"Task {task_id} failed: {error_message}")

def process_extraction(files_data, task_id):
    """Process files extraction with proper sequential ordering and consistent naming - UPDATED for RTF support"""
    try:
        total_files = len(files_data)
        
        # Determine ZIP name and structure based on file count
        if total_files == 1:
            # Single file: Name ZIP after the file, flat structure
            original_filename = files_data[0][0]
            safe_name = secure_filename(original_filename)
            base_name = os.path.splitext(safe_name)[0]
            unique_zip = f"{base_name}_images.zip"
        else:
            # Multiple files: UUID name, folder structure
            unique_zip = f"{uuid.uuid4()}_original_images.zip"
            
        master_zip = os.path.join(app.config['PROCESSED_FOLDER'], unique_zip)

        with tempfile.TemporaryDirectory() as temp_dir:
            has_images = False

            # Save all files to temp dir first
            saved_files = []
            for file_index, (filename, file_content) in enumerate(files_data):
                safe_filename = secure_filename(filename)
                file_path = os.path.join(temp_dir, safe_filename)
                with open(file_path, 'wb') as f:
                    f.write(file_content)
                saved_files.append((safe_filename, file_path))
                update_task_progress(task_id, (file_index / total_files) * 10,
                                   f"Saving file {file_index+1}/{total_files}: {safe_filename}")

            # Prepare a directory for final images to be zipped
            final_images_dir = os.path.join(temp_dir, "final_images")
            os.makedirs(final_images_dir, exist_ok=True)

            # Process each file
            for file_index, (filename, file_path) in enumerate(saved_files):
                update_task_progress(task_id, 10 + (file_index / total_files) * 70,
                                   f"Processing file {file_index+1}/{total_files}: {filename}")

                try:
                    with tempfile.TemporaryDirectory() as file_temp:
                        images = []
                        if filename.lower().endswith('.pdf'):
                            images = extract_images_from_pdf_sequential(file_path, file_temp, task_id, file_index, total_files)
                        elif filename.lower().endswith(('.doc', '.docx', '.rtf')):
                            # Convert DOC/RTF to DOCX if needed
                            final_file_path = file_path
                            
                            if filename.lower().endswith('.doc'):
                                update_task_progress(task_id, 15 + (file_index / total_files) * 10,
                                                   f"Converting DOC {filename} to DOCX format")
                                new_path = doc_to_docx(file_path)
                                if new_path and os.path.exists(new_path):
                                    safe_remove_file(file_path)
                                    final_file_path = new_path
                                else:
                                    continue
                                    
                            elif filename.lower().endswith('.rtf'):
                                update_task_progress(task_id, 15 + (file_index / total_files) * 10,
                                                   f"Converting RTF {filename} to DOCX format")
                                new_path = rtf_to_docx(file_path)
                                if new_path and os.path.exists(new_path):
                                    safe_remove_file(file_path)
                                    final_file_path = new_path
                                else:
                                    continue

                            if final_file_path and os.path.exists(final_file_path):
                                images = extract_images_from_docx_sequential(final_file_path, file_temp, task_id, file_index, total_files)

                        if images:
                            has_images = True
                            doc_name = os.path.splitext(filename)[0]
                            
                            # Determine output folder for this document's images
                            if total_files == 1:
                                # Single file: Images go directly to root of final_images_dir
                                doc_output_dir = final_images_dir
                            else:
                                # Multiple files: Images go to a subfolder named after the document
                                doc_output_dir = os.path.join(final_images_dir, doc_name)
                                os.makedirs(doc_output_dir, exist_ok=True)

                            # Move and rename images to final destination
                            images.sort(key=lambda x: x[2])
                            for img_path, _, counter in images:
                                if os.path.exists(img_path):
                                    # Naming convention: {doc_name}_image_{counter}.jpg
                                    # For single file, this is fine. For multi file, it's also fine inside the folder.
                                    new_name = f"{doc_name}_image_{counter:03d}.jpg"
                                    dest_path = os.path.join(doc_output_dir, new_name)
                                    shutil.move(img_path, dest_path)

                except Exception as e:
                    logger.exception(f"Error processing {filename}: {e}")
                finally:
                    if os.path.exists(file_path):
                        safe_remove_file(file_path)

            # Create final ZIP
            if has_images:
                try:
                    update_task_progress(task_id, 90, "Creating final ZIP archive")
                    
                    # Zip the contents of final_images_dir
                    shutil.make_archive(os.path.splitext(master_zip)[0], 'zip', final_images_dir)
                    
                    cleanup_all_folders()
                    complete_task(task_id, unique_zip)
                except Exception as e:
                    logger.exception(f"Error creating master ZIP: {e}")
                    fail_task(task_id, f"Error creating ZIP: {e}")
            else:
                fail_task(task_id, "No images were extracted from the uploaded files")

    except Exception as e:
        logger.exception(f"Processing error: {e}")
        fail_task(task_id, f"Processing error: {str(e)}")

def process_grayscale_conversion(files_data, task_id):
    """Process grayscale conversion with proper ordering"""
    try:
        unique_zip = f"{uuid.uuid4()}_grayscale_images.zip"
        master_zip = os.path.join(app.config['PROCESSED_FOLDER'], unique_zip)

        with tempfile.TemporaryDirectory() as temp_dir:
            has_images = False
            total_files = len(files_data)

            # Process files in order and assign sequential counters
            saved_files = []
            global_counter = 1
            for file_index, (filename, file_content) in enumerate(files_data):
                safe_filename = secure_filename(filename)
                file_base, file_ext = os.path.splitext(safe_filename)
                temp_file_path = os.path.join(temp_dir, f"temp_{file_index}{file_ext}")

                with open(temp_file_path, 'wb') as f:
                    f.write(file_content)

                update_task(task_id, progress=(file_index / total_files) * 30,
                                   message=f"Processing file {file_index+1}/{total_files}: {safe_filename}")

                # Standardize format with sequential counter
                standardized_path, standardized_name = standardize_image_format(
                    temp_file_path, temp_dir, global_counter, 'JPEG'
                )

                if standardized_path:
                    saved_files.append((standardized_path, standardized_name, global_counter))
                    global_counter += 1
                    has_images = True

                safe_remove_file(temp_file_path)

            # Convert to grayscale
            if saved_files:
                grayscale_images = convert_images_to_grayscale_ordered(saved_files, task_id, 30, 60)
                if grayscale_images:
                    update_task_progress(task_id, 90, "Creating ZIP archive of grayscale images")

                    # Sort by counter before adding to ZIP
                    grayscale_images.sort(key=lambda x: x[2])
                    with zipfile.ZipFile(master_zip, 'w') as zipf:
                        for img_path, img_name, _ in grayscale_images:
                            if os.path.exists(img_path):
                                zipf.write(img_path, img_name)

                    cleanup_all_folders()
                    complete_task(task_id, unique_zip)
                else:
                    fail_task(task_id, "No images were successfully converted")
            else:
                fail_task(task_id, "No valid images found for conversion")

    except Exception as e:
        fail_task(task_id, f"Processing error: {str(e)}")

def process_extract_and_grayscale(files_data, task_id):
    """Process extraction and grayscale conversion with grayscale-specific naming - UPDATED for RTF support"""
    try:
        total_files = len(files_data)
        
        # Determine ZIP name and structure based on file count
        if total_files == 1:
            # Single file: Name ZIP after the file, flat structure
            original_filename = files_data[0][0]
            safe_name = secure_filename(original_filename)
            base_name = os.path.splitext(safe_name)[0]
            unique_zip = f"{base_name}_grayscale.zip"
        else:
            # Multiple files: UUID name, folder structure
            unique_zip = f"{uuid.uuid4()}_grayscale_images.zip"
            
        master_zip = os.path.join(app.config['PROCESSED_FOLDER'], unique_zip)

        with tempfile.TemporaryDirectory() as temp_dir:
            has_images = False

            # Save all files to temp dir first
            saved_files = []
            for file_index, (filename, file_content) in enumerate(files_data):
                safe_filename = secure_filename(filename)
                file_path = os.path.join(temp_dir, safe_filename)
                with open(file_path, 'wb') as f:
                    f.write(file_content)
                saved_files.append((safe_filename, file_path))
                update_task_progress(task_id, (file_index / total_files) * 10,
                                   f"Saving file {file_index+1}/{total_files}: {safe_filename}")

            # Prepare a directory for final images to be zipped
            final_images_dir = os.path.join(temp_dir, "final_images")
            os.makedirs(final_images_dir, exist_ok=True)

            # Process each file
            for file_index, (filename, file_path) in enumerate(saved_files):
                update_task_progress(task_id, 10 + (file_index / total_files) * 70,
                                   f"Processing file {file_index+1}/{total_files}: {filename}")

                try:
                    with tempfile.TemporaryDirectory() as file_temp:
                        images = []
                        if filename.lower().endswith('.pdf'):
                            images = extract_images_from_pdf_sequential(file_path, file_temp, task_id, file_index, total_files)
                        elif filename.lower().endswith(('.doc', '.docx', '.rtf')):
                            # Convert DOC/RTF to DOCX if needed
                            final_file_path = file_path
                            
                            if filename.lower().endswith('.doc'):
                                update_task_progress(task_id, 15 + (file_index / total_files) * 10,
                                                   f"Converting DOC {filename} to DOCX format")
                                new_path = doc_to_docx(file_path)
                                if new_path and os.path.exists(new_path):
                                    safe_remove_file(file_path)
                                    final_file_path = new_path
                                else:
                                    continue
                                    
                            elif filename.lower().endswith('.rtf'):
                                update_task_progress(task_id, 15 + (file_index / total_files) * 10,
                                                   f"Converting RTF {filename} to DOCX format")
                                new_path = rtf_to_docx(file_path)
                                if new_path and os.path.exists(new_path):
                                    safe_remove_file(file_path)
                                    final_file_path = new_path
                                else:
                                    continue

                            if final_file_path and os.path.exists(final_file_path):
                                images = extract_images_from_docx_sequential(final_file_path, file_temp, task_id, file_index, total_files)

                        if images:
                            has_images = True
                            doc_name = os.path.splitext(filename)[0]
                            
                            # Determine output folder for this document's images
                            if total_files == 1:
                                # Single file: Images go directly to root of final_images_dir
                                doc_output_dir = final_images_dir
                            else:
                                # Multiple files: Images go to a subfolder named after the document
                                doc_output_dir = os.path.join(final_images_dir, doc_name)
                                os.makedirs(doc_output_dir, exist_ok=True)

                            # Convert images to grayscale and move to final destination
                            images.sort(key=lambda x: x[2])
                            
                            # Convert to grayscale
                            grayscale_images = convert_images_to_grayscale_ordered(images, task_id, 60, 80)
                            
                            if grayscale_images:
                                for img_path, _, counter in grayscale_images:
                                    if os.path.exists(img_path):
                                        # Naming convention: {doc_name}_image_{counter}.jpg
                                        new_name = f"{doc_name}_image_{counter:03d}.jpg"
                                        dest_path = os.path.join(doc_output_dir, new_name)
                                        shutil.move(img_path, dest_path)

                except Exception as e:
                    logger.exception(f"Error processing {filename}: {e}")
                finally:
                    if os.path.exists(file_path):
                        safe_remove_file(file_path)

            # Create final ZIP
            if has_images:
                try:
                    update_task_progress(task_id, 90, "Creating final ZIP archive")
                    
                    # Zip the contents of final_images_dir
                    shutil.make_archive(os.path.splitext(master_zip)[0], 'zip', final_images_dir)
                    
                    cleanup_all_folders()
                    complete_task(task_id, unique_zip)
                except Exception as e:
                    logger.exception(f"Error creating master ZIP: {e}")
                    fail_task(task_id, f"Error creating ZIP: {e}")
            else:
                fail_task(task_id, "No images were extracted from the uploaded files")

    except Exception as e:
        logger.exception(f"Processing error: {e}")
        fail_task(task_id, f"Processing error: {str(e)}")

def process_document_editing(input_path, output_path):
    """Process DOCX document to remove images and change font to Arial"""
    try:
        # Load the document
        doc = docx.Document(input_path)

        # Remove all images
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Remove all pictures from runs
                if run._element.xpath('.//w:drawing'):
                    run.clear()

        # Remove images from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run._element.xpath('.//w:drawing'):
                                run.clear()

        # Change font to Arial while preserving formatting
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Preserve original formatting
                bold = run.bold
                italic = run.italic
                underline = run.underline
                font_size = Pt(12)
                font_color = run.font.color.rgb if run.font.color.rgb else RGBColor(0, 0, 0)

                # Change font to Arial
                run.font.name = 'Arial'

                # Restore original formatting
                run.bold = bold
                run.italic = italic
                run.underline = underline
                run.font.size = Pt(12)
                run.font.color.rgb = font_color

        # Process text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # Preserve original formatting
                            bold = run.bold
                            italic = run.italic
                            underline = run.underline
                            font_size = run.font.size
                            font_color = run.font.color.rgb if run.font.color.rgb else RGBColor(0, 0, 0)

                            # Change font to Arial
                            run.font.name = 'Arial'

                            # Restore original formatting
                            run.bold = bold
                            run.italic = italic
                            run.underline = underline
                            run.font.size = Pt(12)
                            run.font.color.rgb = font_color

        # Save the processed document
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"Error processing document: {e}")
        return False

def sanitize_filename(name):
    # Hilangkan karakter ilegal
    cleaned = re.sub(r'[<>:"/\\\\|?*]', '', name)
    cleaned = cleaned.strip().rstrip('. ')
    # Pangkas jika terlalu panjang
    return cleaned[:200]

def to_docx(path):
    ext = os.path.splitext(path)[1].lower()
    if ext in ('.doc', '.rtf'):
        if pythoncom is None or win32 is None:
            return path # Return original path if conversion not possible
            
        pythoncom.CoInitialize()
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False; word.DisplayAlerts = False
        abs_in = os.path.abspath(path)
        # Tentukan nama output yang valid
        base = os.path.splitext(abs_in)[0]
        name = os.path.basename(base)
        safe_name = sanitize_filename(name)
        out = os.path.join(os.path.dirname(abs_in), safe_name + '.docx')
        out = os.path.normpath(out)  # gunakan backslashes
        doc = word.Documents.Open(abs_in)
        doc.SaveAs(out, FileFormat=16)
        doc.Close(); word.Quit()
        return out
    return path

def change_font_times_new_roman(docx_path):
    """Ubah semua runs jadi TNR 12pt, pelihara bold/italic/underline/strike/sub/sup."""
    doc = DocxDocument(docx_path)
    # Fungsi bantu untuk satu run
    def apply(run):
        b, i, u = run.bold, run.italic, run.underline
        s, sub, sup = (
            run.font.strike,
            run.font.subscript,
            run.font.superscript
        )
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold, run.italic, run.underline = b, i, u
        run.font.strike, run.font.subscript, run.font.superscript = s, sub, sup

    # Terapkan ke paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            apply(run)
    # Terapkan ke tabel
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply(run)
    doc.save(docx_path)

def merge_with_docxcompose(paths, output_path):
    """
    Gabung dokumen .docx dalam urutan paths,
    tambahkan page break sebelum tiap append.
    """
    master = DocxDocument(paths[0])
    composer = Composer(master)
    for p in paths[1:]:
        master.add_page_break()                   # page break
        composer.append(DocxDocument(p))
    composer.save(output_path)



# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/extract-start', methods=['POST'])
def extract_start():
    files = request.files.getlist('file')
    if not files or files[0].filename == '':
        return "No files selected", 400

    task_id = str(uuid.uuid4())
    save_task(task_id, {
        'status': 'processing',
        'progress': 0,
        'message': 'Starting extraction...',
        'download_filename': None
    })

    files_data = []
    for file in files:
        file_content = file.read()
        files_data.append((file.filename, file_content))

    thread = threading.Thread(target=process_extraction, args=(files_data, task_id))
    thread.daemon = True
    thread.start()

    return task_id

@app.route('/convert-start', methods=['POST'])
def convert_start():
    files = request.files.getlist('images')
    if not files or files[0].filename == '':
        return "No files selected", 400

    task_id = str(uuid.uuid4())
    save_task(task_id, {
        'status': 'processing',
        'progress': 0,
        'message': 'Starting grayscale conversion...',
        'download_filename': None
    })

    files_data = []
    for file in files:
        file_content = file.read()
        files_data.append((file.filename, file_content))

    thread = threading.Thread(target=process_grayscale_conversion, args=(files_data, task_id))
    thread.daemon = True
    thread.start()

    return task_id

@app.route('/extract-grayscale-start', methods=['POST'])
def extract_grayscale_start():
    files = request.files.getlist('file')
    if not files or files[0].filename == '':
        return "No files selected", 400

    task_id = str(uuid.uuid4())
    save_task(task_id, {
        'status': 'processing',
        'progress': 0,
        'message': 'Starting extraction and grayscale conversion...',
        'download_filename': None
    })

    files_data = []
    for file in files:
        file_content = file.read()
        files_data.append((file.filename, file_content))

    thread = threading.Thread(target=process_extract_and_grayscale, args=(files_data, task_id))
    thread.daemon = True
    thread.start()

    return task_id

@app.route('/processing-progress/<task_id>')
def processing_progress(task_id):
    task = get_task(task_id)
    if task:
        if task['status'] == 'completed':
            task['download_url'] = url_for('download_converted_file', filename=task['download_filename'])
        return jsonify(task)
    return jsonify({'status': 'not_found'}), 404

@app.route('/document-edit', methods=['GET', 'POST'])
def document_edit():
    if request.method == 'POST':
        logger.info("=== POST request received for document-edit ===")
        logger.info("request.files: %s", dict(request.files))
        logger.info("request.form: %s", dict(request.form))
        
        # Check if the post request has the file part
        if 'file' not in request.files:
            logger.warning("No 'file' in request.files")
            flash('No file part', 'error')
            return redirect(request.url)

        # Get multiple files
        files = request.files.getlist('file')
        logger.info("Files received: %s", [(f.filename, f.content_length) for f in files])

        # If user does not select file, browser also submit an empty part without filename
        if not files or files[0].filename == '':
            logger.warning("No files selected or empty filename")
            flash('No selected file', 'error')
            return redirect(request.url)

        processed_files = []
        for file in files:
            # Cek apakah file format yang didukung: DOC, DOCX, RTF
            if file and allowed_file(file.filename, {'doc', 'docx', 'rtf'}):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)

                # Konversi file ke DOCX jika diperlukan
                final_file_path = file_path
                original_filename = filename
                
                if filename.lower().endswith('.doc'):
                    print(f"Converting DOC to DOCX: {filename}")
                    new_path = doc_to_docx(file_path)
                    if new_path and os.path.exists(new_path):
                        safe_remove_file(file_path)  # Hapus file DOC asli
                        final_file_path = new_path
                        filename = os.path.basename(new_path)  # Update filename ke DOCX
                    else:
                        flash(f'Error converting DOC file: {original_filename}')
                        safe_remove_file(file_path)
                        continue
                        
                elif filename.lower().endswith('.rtf'):
                    print(f"Converting RTF to DOCX: {filename}")
                    new_path = rtf_to_docx(file_path)
                    if new_path and os.path.exists(new_path):
                        safe_remove_file(file_path)  # Hapus file RTF asli
                        final_file_path = new_path
                        filename = os.path.basename(new_path)  # Update filename ke DOCX
                    else:
                        flash(f'Error converting RTF file: {original_filename}')
                        safe_remove_file(file_path)
                        continue

                # Generate output filename berdasarkan nama file asli
                base_name = os.path.splitext(original_filename)[0]
                output_filename = base_name + '_edited.docx'
                output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)

                # Process the document (sekarang sudah pasti DOCX)
                if process_document_editing(final_file_path, output_path):
                    processed_files.append(output_filename)
                else:
                    flash(f'Error processing file: {original_filename}')

                # Clean up the uploaded/converted file
                if os.path.exists(final_file_path):
                    os.remove(final_file_path)
                    
            else:
                flash('Only .doc, .docx, and .rtf files are allowed')

        if processed_files:
            # Cleanup both folders after processing
            cleanup_all_folders()
            return render_template('document_edit.html', processed=True, filenames=processed_files)
        else:
            return redirect(request.url)
    @after_this_request
    def remove_file(response):
        try:
            file_path = os.path.join(app.config['PROCESSED_FOLDER'], filename)
            time.sleep(0.5)
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Gagal menghapus file: {e}")
        return response

    return send_from_directory(app.config['PROCESSED_FOLDER'], filename, as_attachment=True)

@app.route('/download-doc/<filename>')
def download_doc(filename):
    # Send the processed file
    return_path = send_from_directory(app.config['PROCESSED_FOLDER'], filename, as_attachment=True)

    # Clean up files after sending
    @return_path.call_on_close
    def cleanup():
        # Get the original filename
        original_filename = filename.rsplit('_edited.', 1)[0] + '.docx'

        # Remove the processed file
        processed_path = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        if os.path.exists(processed_path):
            os.remove(processed_path)

        # Remove the original uploaded file if exists
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], original_filename)
        if os.path.exists(upload_path):
            os.remove(upload_path)

    return return_path

@app.route('/grayscale')
def grayscale_form():
    return render_template('grayscale.html')

@app.route('/spine-calculator')
def spine_calculator():
    return render_template('spine_calculator.html')

@app.route('/extract-gray')
def grayscale_extract_form():
    return render_template('warna.html')

@app.route('/merge-docs', methods=['GET','POST'])
def merge_docs_route():
    if request.method == 'POST':
        files = request.files.getlist('file')
        change_font = request.form.get('change_font') == 'on'

        if not files or not files[0].filename:
            flash('❌ Tidak ada file dipilih', 'error')
            return redirect(request.url)

        saved_files = []
        for f in files:
            if allowed_file(f.filename, {'doc', 'docx', 'rtf'}):
                fn = secure_filename(f.filename)
                up = os.path.join(app.config['UPLOAD_FOLDER'], fn)
                f.save(up)
                saved_files.append({'path': up, 'name': f.filename})
        
        if not saved_files:
            flash('❌ Tidak ada file valid yang diunggah (hanya .doc, .docx, .rtf)', 'error')
            return redirect(request.url)

        corrupt_files = []
        valid_docx_paths = []
        temp_files_to_clean = [f['path'] for f in saved_files]

        for file_info in saved_files:
            try:
                # Konversi ke .docx jika perlu
                docx_path = to_docx(file_info['path'])
                if docx_path != file_info['path']:
                    temp_files_to_clean.append(docx_path)

                # Coba buka dokumen untuk memvalidasi
                DocxDocument(docx_path)
                valid_docx_paths.append(docx_path)

            except zipfile.BadZipFile:
                corrupt_files.append(file_info['name'])
                print(f"File korup terdeteksi (BadZipFile): {file_info['name']}")
            except Exception as e:
                corrupt_files.append(file_info['name'])
                print(f"Gagal memproses file {file_info['name']}: {e}")

        if corrupt_files:
            # Hapus semua file temporer jika ada yang korup
            for p in temp_files_to_clean:
                safe_remove_file(p)
            
            flash(Markup(
                "❌ Gagal memproses karena file berikut rusak atau tidak dapat dibaca:<br> - " + 
                "<br> - ".join(corrupt_files) +
                "<br><br>Silakan perbaiki file tersebut dan coba lagi."
            ), 'error')
            return redirect(request.url)

        if not valid_docx_paths:
            flash('❌ Tidak ada file valid yang bisa diproses.', 'error')
            # Hapus semua file temporer jika tidak ada yang valid
            for p in temp_files_to_clean:
                safe_remove_file(p)
            return redirect(request.url)

        # Lanjutkan merge jika semua file valid
        try:
            out_name = f"merged_{uuid.uuid4().hex}.docx"
            out_path = os.path.join(app.config['PROCESSED_FOLDER'], out_name)

            master = DocxDocument(valid_docx_paths[0])
            composer = Composer(master)
            for p in valid_docx_paths[1:]:
                master.add_page_break()
                composer.append(DocxDocument(p))
            composer.save(out_path)

            if change_font:
                change_font_times_new_roman(out_path)

            download_url = url_for('download_file', filename=out_name)
            flash(Markup(
                "✅ Dokumen berhasil digabung!<br><br>"
                f"<a href='{download_url}' class='btn'>💾 Download di sini</a>"
            ))
        
        except Exception as e:
            flash(f'❌ Terjadi error saat menggabungkan dokumen: {e}', 'error')
        
        finally:
            # Hapus semua file temporer setelah selesai
            for p in temp_files_to_clean:
                safe_remove_file(p)

        return redirect(request.url)

    return render_template('merge_docs.html')

@app.route('/image-upscaler')
def image_upscaler_form():
    return render_template('image_upscaler.html')

@app.route('/process-image', methods=['POST'])
def process_image():
    if 'image' not in request.files:
        return jsonify({'success': False, 'message': 'No image file provided'}), 400

    file = request.files['image']
    operation = request.form.get('operation')

    if file.filename == '':
        return jsonify({'success': False, 'message': 'No selected image'}), 400

    if file and allowed_file(file.filename, {'jpg', 'jpeg', 'png'}):
        unique_id = str(uuid.uuid4())
        original_filename = secure_filename(file.filename)
        temp_input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{original_filename}")
        file.save(temp_input_path)

        if operation == 'upscale':
            try:
                task_id, scale_value = upload_image(temp_input_path)
                if task_id:
                    download_url = check_status(task_id, scale_value)
                    if download_url:
                        # Generate a unique filename for the processed image
                        output_filename = f"{unique_id}_upscaled_{original_filename}"
                        output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
                        if download_result(download_url, output_path):
                            # Cleanup temporary uploaded file
                            safe_remove_file(temp_input_path)
                            # Return the URL to the processed image
                            return jsonify({'success': True, 'image_url': url_for('download_file', filename=output_filename)}), 200
                        else:
                            safe_remove_file(temp_input_path)
                            return jsonify({'success': False, 'message': 'Failed to download upscaled image'}), 500
                    else:
                        safe_remove_file(temp_input_path)
                        return jsonify({'success': False, 'message': 'Upscaling process failed or timed out'}), 500
                else:
                    safe_remove_file(temp_input_path)
                    return jsonify({'success': False, 'message': 'Failed to upload image for upscaling'}), 500
            except Exception as e:
                safe_remove_file(temp_input_path)
                print(f"Error during upscaling: {e}")
                return jsonify({'success': False, 'message': f'An error occurred during upscaling: {str(e)}'}), 500
        else:
            safe_remove_file(temp_input_path)
            return jsonify({'success': False, 'message': f'Operation {operation} not supported yet.'}), 400
    else:
        return jsonify({'success': False, 'message': 'Invalid file type. Only JPG, JPEG, PNG are allowed.'}), 400


# Base directory untuk folder berkas
BERKAS_DIR = os.path.join(os.getcwd(), 'files')

def format_file_size(size):
    """Format ukuran file ke format yang mudah dibaca"""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size < 1024.0:
            return f"{size:.1f} {unit}"
        size /= 1024.0
    return f"{size:.1f} TB"

@app.route('/berkas/search')
def search_berkas():
    query = request.args.get('q', '').lower()
    path = request.args.get('path', '')
    
    if not query:
        return jsonify({'results': []})

    search_root = os.path.join(BERKAS_DIR, path)
    
    # Security check
    if not os.path.abspath(search_root).startswith(os.path.abspath(BERKAS_DIR)):
        return jsonify({'error': 'Access denied'}), 403
        
    results = []
    
    try:
        for root, dirs, files in os.walk(search_root):
            # Calculate relative path from search_root for display
            rel_path = os.path.relpath(root, BERKAS_DIR).replace('\\', '/')
            if rel_path == '.': rel_path = ''
            
            # Search in directories
            for dirname in dirs:
                if query in dirname.lower():
                    full_path = os.path.join(root, dirname)
                    rel_full_path = os.path.join(rel_path, dirname).replace('\\', '/')
                    
                    results.append({
                        'name': dirname,
                        'type': 'folder',
                        'path': rel_full_path,
                        'parent': rel_path if rel_path else '.',
                        'modified': datetime.fromtimestamp(os.path.getmtime(full_path)).strftime('%Y-%m-%d %H:%M'),
                        'size': '-'
                    })
            
            # Search in files
            for filename in files:
                if query in filename.lower():
                    full_path = os.path.join(root, filename)
                    rel_full_path = os.path.join(rel_path, filename).replace('\\', '/')
                    stat = os.stat(full_path)
                    
                    results.append({
                        'name': filename,
                        'type': 'file',
                        'path': rel_full_path,
                        'parent': rel_path if rel_path else '.',
                        'modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M'),
                        'size': format_file_size(stat.st_size)
                    })
                    
            # Limit results to prevent overload
            if len(results) > 100:
                break
                
    except Exception as e:
        logger.error(f"Search error: {e}")
        return jsonify({'error': str(e)}), 500
        
    return jsonify({'results': results})

@app.route('/berkas/', defaults={'path': ''})
@app.route('/berkas/<path:path>')
def serve_berkas_files(path):
    full_path = os.path.join(BERKAS_DIR, path)
    
    # Cek keamanan path (prevent directory traversal)
    if not os.path.abspath(full_path).startswith(os.path.abspath(BERKAS_DIR)):
        abort(403)
    
    # Jika directory, tampilkan listing
    if os.path.isdir(full_path):
        files = []
        
        # Dapatkan informasi file
        for filename in os.listdir(full_path):
            file_path = os.path.join(full_path, filename)
            # Gunakan forward slash untuk URL path
            file_url_path = os.path.join(path, filename).replace('\\', '/') if path else filename
            
            # Dapatkan informasi file
            try:
                stat = os.stat(file_path)
                is_dir = os.path.isdir(file_path)
                
                files.append({
                    'name': filename,
                    'is_dir': is_dir,
                    'size': 0 if is_dir else stat.st_size,
                    'size_formatted': '-' if is_dir else format_file_size(stat.st_size),
                    'modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M'),
                    'url_path': file_url_path
                })
            except OSError:
                continue
        
        # Sort: folders first, then files (alphabetically)
        files.sort(key=lambda x: (not x['is_dir'], x['name'].lower()))
        
        # Prepare breadcrumbs
        breadcrumbs = []
        if path:
            parts = path.replace('\\', '/').split('/')
            current_crumb = ''
            for part in parts:
                if part:
                    current_crumb = f"{current_crumb}/{part}" if current_crumb else part
                    breadcrumbs.append({
                        'name': part,
                        'url': f"/berkas/{current_crumb}"
                    })

        # Prepare parent URL
        parent_url = '/berkas/'
        if path:
            parent_path = os.path.dirname(path.rstrip('/'))
            if parent_path:
                parent_url = f"/berkas/{parent_path.replace('\\', '/')}"

        # Calculate stats
        folder_count = sum(1 for f in files if f['is_dir'])
        file_count = len(files) - folder_count
        total_size = sum(f['size'] for f in files if not f['is_dir'])
        total_size_formatted = format_file_size(total_size)

        return render_template('berkas_browser.html',
                             files=files,
                             current_path=path,
                             breadcrumbs=breadcrumbs,
                             parent_url=parent_url,
                             folder_count=folder_count,
                             file_count=file_count,
                             total_size_formatted=total_size_formatted)

    # Jika file, serve file tersebut
    elif os.path.isfile(full_path):
        return send_file(full_path)
    
    # File tidak ditemukan
    else:
        abort(404)

from flask import session

@app.route('/batch-converter', methods=['GET', 'POST'])
def batch_converter():
    if request.method == 'POST':
        files = request.files.getlist('image_files')
        output_format = request.form.get('output_format')

        if not files or not any(f.filename for f in files):
            flash('No files selected', 'error')
            return redirect(request.url)

        session_id = str(uuid.uuid4())
        session_folder = os.path.join(app.config['PROCESSED_FOLDER'], session_id)
        os.makedirs(session_folder, exist_ok=True)

        converted_files = []
        for file in files:
            if file and allowed_file(file.filename, {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp', 'tiff'}):
                try:
                    filename = secure_filename(file.filename)
                    base_filename = os.path.splitext(filename)[0]
                    new_filename = f"{base_filename}.{output_format.lower()}"
                    output_path = os.path.join(session_folder, new_filename)

                    image = Image.open(file.stream)

                    # Handle RGBA to RGB conversion for JPG and BMP
                    if output_format.upper() in ['JPG', 'JPEG', 'BMP'] and image.mode == 'RGBA':
                        image = image.convert('RGB')
                    
                    save_options = {}
                    if output_format.upper() == 'JPG' or output_format.upper() == 'JPEG':
                        save_options = {'quality': 100, 'optimize': True, 'subsampling': 0}
                    elif output_format.upper() == 'WEBP':
                        save_options = {'quality': 100, 'lossless': True}
                    elif output_format.upper() == 'PNG':
                        save_options = {'optimize': True}
                    elif output_format.upper() == 'TIFF':
                        save_options = {'compression': 'tiff_lzw'}


                    image.save(output_path, **save_options)
                    converted_files.append(new_filename)
                except Exception as e:
                    flash(f"Error converting {file.filename}: {e}", 'error')

        if not converted_files:
            flash('No files were converted.', 'error')
            return redirect(url_for('batch_converter'))

        session['converted_files'] = converted_files
        session['session_folder_id'] = session_id
        return redirect(url_for('converter_results'))

    return render_template('batch_converter.html')

@app.route('/converter-results')
def converter_results():
    converted_files = session.get('converted_files', [])
    return render_template('converter_result.html', converted_files=converted_files)

@app.route('/download-converted/<filename>')
def download_converted_file(filename):
    session_folder_id = session.get('session_folder_id')
    if not session_folder_id:
        abort(404)
    
    session_folder = os.path.join(app.config['PROCESSED_FOLDER'], session_folder_id)
    return send_from_directory(session_folder, filename, as_attachment=True)

@app.route('/download-zip')
def download_zip():
    session_folder_id = session.get('session_folder_id')
    converted_files = session.get('converted_files', [])

    if not session_folder_id or not converted_files:
        abort(404)

    session_folder = os.path.join(app.config['PROCESSED_FOLDER'], session_folder_id)
    zip_path = os.path.join(app.config['PROCESSED_FOLDER'], f"{session_folder_id}.zip")

    with zipfile.ZipFile(zip_path, 'w') as zf:
        for filename in converted_files:
            file_path = os.path.join(session_folder, filename)
            if os.path.exists(file_path):
                zf.write(file_path, arcname=filename)

    @after_this_request
    def cleanup(response):
        safe_remove_file(zip_path)
        shutil.rmtree(session_folder, ignore_errors=True)
        session.pop('converted_files', None)
        session.pop('session_folder_id', None)
        return response

    return send_file(zip_path, as_attachment=True, download_name='converted_images.zip')


if __name__ == '__main__':
    # Start cleanup thread
    cleanup_thread = threading.Thread(target=periodic_cleanup)
    cleanup_thread.daemon = True
    cleanup_thread.start()
    print("✅ Cleanup berkala dimulai (setiap 30 menit)")
    
    # Start Flask app
    print("🚀 Starting Image Processor Application...")
    print("📁 Maksimal file di uploads: 500")
    print("📁 Maksimal file di processed: 5")
    app.run(host='0.0.0.0', port=5000, debug=True)
